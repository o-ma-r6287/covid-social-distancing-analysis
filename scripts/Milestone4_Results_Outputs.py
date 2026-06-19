import os
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import statsmodels.formula.api as smf
from scipy.stats import chi2_contingency

data_path = "/Users/omara-rahman/Desktop/Semester 2 MPH/Computer based class/Python BIOS 584/Final_project/dataverse_files 2/cleanedCDF.csv"
output_folder = "/Users/omara-rahman/Desktop/Semester 2 MPH/Computer based class/Python BIOS 584/Final_project/dataverse_files 2/"

#Loaded Data
df = pd.read_csv(data_path)

#Theme
navy = "#0A2342"
gold = "#D4AF37"

#RQ1 Linear Regression
linear_iso = smf.ols(
    "soc_dist6 ~ age + C(gender_group) + C(education_group) + C(employment_group) + C(income_group) + C(ideo_group) + C(party_group) + C(marital_group)",
    data=df
).fit()

#PLOT
linear_ci = linear_iso.conf_int()

plt.figure(figsize=(10, 6))
plt.errorbar(
    linear_iso.params.values,
    linear_iso.params.index,
    xerr=[
        linear_iso.params.values - linear_ci[0].values,
        linear_ci[1].values - linear_iso.params.values
    ],
    fmt="o",
    color=navy
)

plt.axvline(x=0, color=gold)
plt.title("RQ1: Linear Regression Predictors of Self-Isolation Score")
plt.xlabel("Coefficient")
plt.ylabel("Predictor")
plt.tight_layout()
plt.savefig(output_folder + "NEW_RQ1_plot.png", dpi=300)
plt.close()

#RQ2 Logistic Regression
logit_df = df[df["gender_group"] != "Other"].copy()

logit_mask = smf.logit(
    "mask_binary ~ age + C(gender_group) + C(education_group) + C(ideo_group) + C(party_group)",
    data=logit_df
).fit()

#PLOT
logit_ci = logit_mask.conf_int()

plt.figure(figsize=(10, 6))
plt.errorbar(
    np.exp(logit_mask.params.values),
    logit_mask.params.index,
    xerr=[
        np.exp(logit_mask.params.values) - np.exp(logit_ci[0].values),
        np.exp(logit_ci[1].values) - np.exp(logit_mask.params.values)
    ],
    fmt="o",
    color=navy
)

plt.axvline(x=1, color=gold)
plt.title("RQ2: Logistic Regression Predictors of Mask Willingness")
plt.xlabel("Odds Ratio")
plt.ylabel("Predictor")
plt.tight_layout()
plt.savefig(output_folder + "NEW_RQ2_plot.png", dpi=300)
plt.close()

#RQ3 Chi-Square
chi_table = pd.crosstab(df["party_group"], df["mask_binary"])
chi2, p, dof, expected = chi2_contingency(chi_table)

#PLOT
chi_percent = pd.crosstab(
    df["party_group"],
    df["mask_binary"],
    normalize="index"
) * 100

chi_percent.plot(
    kind="bar",
    stacked=True,
    color=[gold, navy],
    figsize=(10, 6)
)

plt.title("RQ3: Mask Willingness by Political Party")
plt.xlabel("Political Party")
plt.ylabel("Percent")
plt.tight_layout()
plt.savefig(output_folder + "NEW_RQ3_plot.png", dpi=300)
plt.close()

#Clean names
def clean_names(name):
    if name == "Intercept":
        return "Intercept"
    if name == "age":
        return "Age"

    name = name.replace("C(", "").replace(")", "")
    name = name.replace("[T.", " (").replace("]", ")")
    name = name.replace("gender_group", "Gender")
    name = name.replace("education_group", "Education")
    name = name.replace("employment_group", "Employment")
    name = name.replace("income_group", "Income")
    name = name.replace("ideo_group", "Ideology")
    name = name.replace("party_group", "Party")
    name = name.replace("marital_group", "Marital Status")
    name = name.replace("_", " ")
    return name

#Format CI clean
def make_ci_text(point, lower, upper):
    point = round(point, 3)
    lower = round(lower, 3)
    upper = round(upper, 3)
    return f"{point} [{lower}, {upper}]"

#RQ1 Final Table
rq1_rows = []
ci = linear_iso.conf_int()

for pred in linear_iso.params.index:
    rq1_rows.append({
        "Predictor": clean_names(pred),
        "Slope [95% CI]": make_ci_text(
            linear_iso.params[pred],
            ci.loc[pred, 0],
            ci.loc[pred, 1]
        ),
        "p-value": 0.0001 if linear_iso.pvalues[pred] < 0.0001 else round(linear_iso.pvalues[pred], 4)
    })

rq1_final = pd.DataFrame(rq1_rows)
rq1_final.to_csv(output_folder + "NEW_FINAL_RQ1.csv", index=False)

#RQ2 Final Table
rq2_rows = []
ci = logit_mask.conf_int()

for pred in logit_mask.params.index:
    rq2_rows.append({
        "Predictor": clean_names(pred),
        "Odds Ratio [95% CI]": make_ci_text(
            np.exp(logit_mask.params[pred]),
            np.exp(ci.loc[pred, 0]),
            np.exp(ci.loc[pred, 1])
        ),
        "p-value": 0.0001 if logit_mask.pvalues[pred] < 0.0001 else round(logit_mask.pvalues[pred], 4)
    })

rq2_final = pd.DataFrame(rq2_rows)
rq2_final.to_csv(output_folder + "NEW_FINAL_RQ2.csv", index=False)

#RQ3 Final Table
rq3_final = pd.DataFrame({
    "Test": ["Party × Mask Willingness"],
    "Chi-square": [round(chi2, 3)],
    "df": [dof],
    "p-value": 0.0001 if p < 0.0001 else round(p, 4)
})

rq3_final.to_csv(output_folder + "NEW_FINAL_RQ3.csv", index=False)

#All RQ1, RQ2, and RQ3 result tables and plots


from docx import Document
#Create document
doc = Document()

#Add title
doc.add_heading('Final Results Tables', level=1)


#Function to add dataframe to Word
def add_table_to_doc(document, df, title):
    document.add_heading(title, level=2)

    table = document.add_table(rows=1, cols=len(df.columns))

    #Header row
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)

    #Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


#Add your tables
add_table_to_doc(doc, rq1_final, "Table 1. Linear Regression (Self-Isolation)")
add_table_to_doc(doc, rq2_final, "Table 2. Logistic Regression (Mask Willingness)")
add_table_to_doc(doc, rq3_final, "Table 3. Chi-Square Results")

#Save document
doc.save(output_folder + "FINAL_RESULTS_TABLES.docx")