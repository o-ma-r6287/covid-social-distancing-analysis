from docx import Document

# Create document
doc = Document()

# Add title
doc.add_heading('Final Results Tables', level=1)


# Function to add dataframe to Word
def add_table_to_doc(document, df, title):
    document.add_heading(title, level=2)

    table = document.add_table(rows=1, cols=len(df.columns))

    # Header row
    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = str(col)

    # Data rows
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)


# Add your tables
add_table_to_doc(doc, rq1_final, "Table 1. Linear Regression (Self-Isolation)")
add_table_to_doc(doc, rq2_final, "Table 2. Logistic Regression (Mask Willingness)")
add_table_to_doc(doc, rq3_final, "Table 3. Chi-Square Results")

# Save document
doc.save(output_folder + "FINAL_RESULTS_TABLES.docx")